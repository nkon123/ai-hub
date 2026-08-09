"""PDF/Word 확장 Loader (04-knowledge-platform.md §2.3) tests.

Two different guarantees are tested here, and they behave very differently
on THIS development machine (no network installs allowed — see
open-decisions.md D-073) vs. the target Windows PC (packages installable):

1. "Missing dependency degrades to a clear Korean error, not a crash" —
   `pypdf`/`python-docx` are in fact NOT installed in this dev environment
   (see indexing-runtime's pyproject.toml: they're declared but the lockfile
   has never been resolved against a network). That makes THIS machine the
   perfect place to prove the degrade path for real, with no mocking.

2. "Real content extraction works" — these tests need the actual library
   importable, which is only true once `uv sync` runs on a machine with
   network/mirror access. They are marked `pytest.importorskip(...)` with an
   explicit reason naming the missing package, per this task's instruction
   to never fake a passing test. Someone running this suite after `uv sync`
   on the target Windows PC is the one who will actually exercise them.
"""

from __future__ import annotations

import pytest
from indexing_runtime.loaders import MissingLoaderDependencyError, load_document


def test_load_pdf_actually_missing_raises_missing_loader_dependency_error(tmp_path):
    # No importorskip here on purpose — this asserts the *absence* case.
    try:
        import pypdf  # noqa: F401

        pytest.skip("pypdf is installed in this environment; degrade path not exercised")
    except ImportError:
        pass

    path = tmp_path / "a.pdf"
    path.write_bytes(b"%PDF-1.4\n%%EOF")  # content is irrelevant; import fails first

    with pytest.raises(MissingLoaderDependencyError, match="pypdf"):
        load_document(path, document_id="kb-1:a.pdf")


def test_load_docx_actually_missing_raises_missing_loader_dependency_error(tmp_path):
    try:
        import docx  # noqa: F401

        pytest.skip("python-docx is installed in this environment; degrade path not exercised")
    except ImportError:
        pass

    path = tmp_path / "a.docx"
    path.write_bytes(b"PK\x03\x04")  # content is irrelevant; import fails first

    with pytest.raises(MissingLoaderDependencyError, match="python-docx"):
        load_document(path, document_id="kb-1:a.docx")


def _make_minimal_pdf(text: str) -> bytes:
    """Hand-built minimal single-page PDF with one Tj text-showing operator.

    NOTE: byte offsets in the xref table are computed programmatically below
    (never hardcoded), which removes the usual "miscounted bytes" failure
    mode of hand-written PDF fixtures. The object/stream syntax itself
    follows the standard minimal-PDF skeleton (Catalog -> Pages -> Page ->
    Font + content stream) that pypdf's own test fixtures use the same
    shape of. That said, this fixture has NOT been executed against a real
    `pypdf.PdfReader` on this machine (pypdf isn't installed here — see
    module docstring) — if `test_load_pdf_text_extracts_visible_text` fails
    once run on a machine with pypdf installed, check this builder first
    before assuming a regression in `pdf_loader.load_pdf_text`.
    """
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /Resources << /Font << /F1 4 0 R >> >> "
        b"/MediaBox [0 0 200 200] /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    stream_content = f"BT /F1 12 Tf 10 100 Td ({text}) Tj ET".encode("latin-1")
    objects.append(
        b"<< /Length "
        + str(len(stream_content)).encode("ascii")
        + b" >>\nstream\n"
        + stream_content
        + b"\nendstream"
    )

    buf = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for i, obj in enumerate(objects, start=1):
        offsets.append(len(buf))
        buf += f"{i} 0 obj\n".encode("ascii")
        buf += obj
        buf += b"\nendobj\n"

    xref_offset = len(buf)
    count = len(objects) + 1  # +1 for the free-list head entry (object 0)
    buf += f"xref\n0 {count}\n".encode("ascii")
    buf += b"0000000000 65535 f \n"
    for off in offsets:
        buf += f"{off:010d} 00000 n \n".encode("ascii")
    buf += (
        f"trailer\n<< /Size {count} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF"
    ).encode("ascii")
    return bytes(buf)


def test_load_pdf_text_extracts_visible_text(tmp_path):
    pytest.importorskip("pypdf", reason="pypdf가 설치되어 있지 않습니다 (uv sync 필요)")

    path = tmp_path / "a.pdf"
    path.write_bytes(_make_minimal_pdf("Hello World"))

    doc = load_document(path, document_id="kb-1:a.pdf")

    assert "Hello World" in doc["content"]
    assert doc["metadata"]["mime_type"] == "application/pdf"
    assert doc["metadata"]["file_name"] == "a.pdf"
    assert doc["metadata"]["title"] == "a"


def test_load_docx_text_extracts_paragraphs(tmp_path):
    docx = pytest.importorskip("docx", reason="python-docx가 설치되어 있지 않습니다 (uv sync 필요)")

    path = tmp_path / "a.docx"
    document = docx.Document()
    document.add_paragraph("첫 번째 문단")
    document.add_paragraph("두 번째 문단")
    document.save(path)

    doc = load_document(path, document_id="kb-1:a.docx")

    assert doc["content"] == "첫 번째 문단\n두 번째 문단"
    assert doc["metadata"]["mime_type"] == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert doc["metadata"]["file_name"] == "a.docx"
