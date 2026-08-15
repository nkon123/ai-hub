"""POST /indexing/v1/extract-text — server-side text extraction for the P12
AI 추천 button's .pdf/.docx path (portal-api relays here; see
`apps/portal-api/src/portal_api/routers/knowledge_text_extract.py`).

This dev machine deliberately does NOT have `pypdf`/`python-docx` installed
(open-decisions.md D-073, company policy — no network/binary installs on
this hardware), which makes it the right place to prove the
MissingLoaderDependencyError degrade path for real (no mocking), exactly the
same reasoning `test_pdf_docx_loaders.py` already documents. The "real
content extraction" happy path for .pdf/.docx is `pytest.importorskip`'d and
is NOT exercised on this machine — see that file's docstring.
"""

from __future__ import annotations

import io

import pytest
from fastapi.testclient import TestClient
from indexing_runtime.main import app
from indexing_runtime.settings import EXTRACT_TEXT_MAX_UPLOAD_BYTES

client = TestClient(app)


def _upload(filename: str, content: bytes, content_type: str = "application/octet-stream"):
    return client.post(
        "/indexing/v1/extract-text",
        files={"file": (filename, io.BytesIO(content), content_type)},
    )


# --- 지원 형식 (텍스트) ---


def test_extract_text_supported_markdown_returns_excerpt() -> None:
    resp = _upload("notes.md", "# 제목\n\n본문 내용입니다.".encode())
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["excerpt"] == "# 제목\n\n본문 내용입니다."
    assert body["trace_id"]


def test_extract_text_bounds_excerpt_to_setting(monkeypatch: pytest.MonkeyPatch) -> None:
    import indexing_runtime.main as main_module

    monkeypatch.setattr(main_module, "EXTRACT_TEXT_EXCERPT_MAX_CHARS", 5)
    resp = _upload("notes.txt", b"0123456789")
    assert resp.status_code == 200, resp.text
    assert resp.json()["excerpt"] == "01234"


# --- 미지원 형식 (VALIDATION_ERROR, 422) — "이 형식은 추천을 지원하지 않습니다" ---


def test_extract_text_unsupported_format_returns_422_validation_error() -> None:
    resp = _upload("spreadsheet.xlsx", b"not a real xlsx")
    assert resp.status_code == 422, resp.text
    body = resp.json()
    assert body["error"]["code"] == "VALIDATION_ERROR"
    assert "지원하지 않습니다" in body["error"]["message"]
    assert body["error"]["trace_id"]


# --- 의존성 미설치 (DEPENDENCY_UNAVAILABLE, 503) — "서버에 ... 추출 의존성이 설치되어
#     있지 않습니다" — 미지원 형식과 반드시 다른 코드/문구여야 한다 ---
#
# 아래 두 테스트는 실제로 pypdf/python-docx가 없을 때만 그 경로를 때린다(이
# 개발 머신에 실제로 없다면 그대로 실증, 있다면 skip — test_pdf_docx_loaders.py
# 와 동일한 관례). 그와 별개로, 이 배포 환경과 무관하게 그 경로를 항상 결정적
# (deterministic)으로 검증하기 위해 `load_text_from_bytes`를 직접
# monkeypatch하는 테스트를 하나 더 둔다(바로 아래) — 두 테스트가 서로를
# 대체하지 않는다.


def test_extract_text_pdf_missing_dependency_returns_503_dependency_unavailable() -> None:
    try:
        import pypdf  # noqa: F401

        pytest.skip("pypdf is installed in this environment; degrade path not exercised")
    except ImportError:
        pass

    resp = _upload("report.pdf", b"%PDF-1.4\n%%EOF")
    assert resp.status_code == 503, resp.text
    body = resp.json()
    assert body["error"]["code"] == "DEPENDENCY_UNAVAILABLE"
    assert "pypdf" in body["error"]["message"]
    assert body["error"]["trace_id"]
    # Must NOT be the same message as the unsupported-format case — an
    # operator reading this needs to know installing a dependency would fix
    # it, which "이 형식은 지원하지 않습니다" would falsely deny.
    assert "지원하지 않습니다" not in body["error"]["message"]


def test_extract_text_docx_missing_dependency_returns_503_dependency_unavailable() -> None:
    try:
        import docx  # noqa: F401

        pytest.skip("python-docx is installed in this environment; degrade path not exercised")
    except ImportError:
        pass

    resp = _upload(
        "memo.docx",
        b"PK\x03\x04",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert resp.status_code == 503, resp.text
    body = resp.json()
    assert body["error"]["code"] == "DEPENDENCY_UNAVAILABLE"
    assert "python-docx" in body["error"]["message"]


def test_extract_text_dependency_unavailable_path_is_deterministic(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Exercises the DEPENDENCY_UNAVAILABLE branch regardless of whether
    pypdf/python-docx happen to be installed on the machine running this
    suite (unlike the two tests above, which only fire when the dependency
    is actually absent) — forces `load_text_from_bytes` to raise, then
    asserts the endpoint's degrade behavior: 503, the distinguishing error
    code, and the exact exception message propagated verbatim (never
    swallowed/replaced with a generic string)."""
    import indexing_runtime.loaders as loaders_module
    import indexing_runtime.main as main_module

    def _raise_missing_dependency(raw: bytes, filename: str) -> str:
        raise loaders_module.MissingLoaderDependencyError(
            "PDF 로더에 필요한 pypdf가 설치되어 있지 않습니다. "
            "indexing-runtime 패키지를 다시 설치한 뒤 다시 시도하세요 (uv sync)."
        )

    monkeypatch.setattr(main_module, "load_text_from_bytes", _raise_missing_dependency)

    resp = _upload("report.pdf", b"%PDF-1.4\n%%EOF")
    assert resp.status_code == 503, resp.text
    body = resp.json()
    assert body["error"]["code"] == "DEPENDENCY_UNAVAILABLE"
    assert "pypdf" in body["error"]["message"]
    assert "지원하지 않습니다" not in body["error"]["message"]
    assert body["error"]["trace_id"]


# --- 실제 .docx 추출 (이 머신에 python-docx가 설치되어 있을 때만 실행) ---


def test_extract_text_docx_real_content_extraction() -> None:
    docx = pytest.importorskip("docx", reason="python-docx가 설치되어 있지 않습니다 (uv sync 필요)")

    buf = io.BytesIO()
    document = docx.Document()
    document.add_paragraph("첫 번째 문단")
    document.add_paragraph("두 번째 문단")
    document.save(buf)

    resp = _upload(
        "report.docx",
        buf.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["excerpt"] == "첫 번째 문단\n두 번째 문단"
    assert body["trace_id"]


# --- 크기 초과 (VALIDATION_ERROR, 413) ---


def test_extract_text_oversized_upload_rejected_before_extraction(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    import indexing_runtime.main as main_module

    monkeypatch.setattr(main_module, "EXTRACT_TEXT_MAX_UPLOAD_BYTES", 10)
    resp = _upload("big.txt", b"0123456789ABCDEF")  # 16 bytes > 10-byte cap
    assert resp.status_code == 413, resp.text
    body = resp.json()
    assert body["error"]["code"] == "VALIDATION_ERROR"
    assert "너무 큽니다" in body["error"]["message"]


def test_extract_text_max_upload_bytes_setting_is_a_real_bound() -> None:
    # Sanity: the setting exists and is a positive int (not accidentally 0
    # or a str), i.e. it actually bounds something.
    assert isinstance(EXTRACT_TEXT_MAX_UPLOAD_BYTES, int)
    assert EXTRACT_TEXT_MAX_UPLOAD_BYTES > 0


# --- 빈 추출 결과 (VALIDATION_ERROR, 422) ---


def test_extract_text_empty_content_returns_validation_error() -> None:
    resp = _upload("empty.txt", b"   \n\n  ")
    assert resp.status_code == 422, resp.text
    body = resp.json()
    assert body["error"]["code"] == "VALIDATION_ERROR"
    assert "비어" in body["error"]["message"]
